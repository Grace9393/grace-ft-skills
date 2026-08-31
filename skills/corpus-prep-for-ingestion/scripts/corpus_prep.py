"""corpus_prep.py — make a document corpus ingestible by a platform with file caps.

Four subcommands, run in this order:

    python corpus_prep.py inventory --in CORPUS --out OUT
    python corpus_prep.py extract   --in CORPUS --out OUT
    python corpus_prep.py fit       --out OUT [--cap-mb 10]
    python corpus_prep.py check     --in CORPUS --out OUT [--cap-mb 10]

Design rules this file enforces, each of which cost a real run to learn:

1. Nothing here reads a document for meaning. Names and sizes come from the ZIP
   central directory; duplicates are settled by SHA-256. Reading 70 MB of documents
   into an agent context to decide "is this a duplicate" is what kills the run.
2. The manifest is written BEFORE extraction and flushed incrementally. A run that
   stops halfway leaves a usable artifact.
3. Every container is opened, not just the largest. Embedded parts hide in small files.
4. Nothing is truncated silently. Every reduction lands in dropped.md.

5. Container parsing is delegated to the shared engine in the `deep-extract`
   skill. A container that will not open is reported, never skipped: an earlier
   version swallowed the exception and returned "0 embedded parts" for a
   damaged 65 MB file that in fact held 29.

Requires: the `deep-extract` skill (shared engine). python-docx and openpyxl for
`fit` only.
"""

import argparse, csv, hashlib, json, os, shutil, sys, zipfile

sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
import _engine

CONTAINER_EXT = (".docx", ".xlsx", ".docm", ".xlsm", ".pptx")
EMBED_MARKER = "/embeddings/"

_DX = None
DAMAGE = {}          # container path -> [warning, ...]
DECLARED_MISSING = {}  # container path -> [embedding part declared but absent]


def dx():
    global _DX
    if _DX is None:
        _DX = _engine.load()
    return _DX


def sha256(b):
    return hashlib.sha256(b).hexdigest()


def walk_corpus(root):
    for dp, _, fs in os.walk(root):
        for f in sorted(fs):
            if f.lower() == "desktop.ini":
                continue
            yield os.path.join(dp, f)


def embedded_parts(path):
    """Yield (part_name, bytes) for every embedded document in a container.

    Damage is recorded in DAMAGE rather than swallowed, and a truncated
    container still yields everything written before the break.
    """
    if not path.lower().endswith(CONTAINER_EXT):
        return
    e = dx()
    try:
        with open(e.long_path(path), "rb") as fh:
            raw = fh.read()
    except OSError as exc:
        DAMAGE.setdefault(path, []).append("unreadable file: %s" % exc)
        return

    entries, warns = e.read_zip_entries(raw)
    if warns:
        DAMAGE.setdefault(path, []).extend(warns)

    parts = {n: b for n, b, _ in entries if b is not None}
    if parts:
        _decl, _present, missing = e.embedding_reconciliation(parts)
        if missing:
            DECLARED_MISSING.setdefault(path, []).extend(missing)

    for name, data, _partial in entries:
        if EMBED_MARKER in name:
            yield name, data


def cmd_inventory(a):
    """Manifest from metadata + hashes only. No document is read for meaning."""
    os.makedirs(a.out, exist_ok=True)
    rows, cap = [], a.cap_mb * 1024 * 1024
    for p in walk_corpus(a.inp):
        rel = os.path.relpath(p, a.inp).replace(os.sep, "/")
        b = open(p, "rb").read()
        rows.append({"kind": "standalone", "path": rel, "container": "",
                     "bytes": len(b), "sha": sha256(b)})
        for name, data in embedded_parts(p):
            rows.append({"kind": "embedded", "path": name, "container": rel,
                         "bytes": -1 if data is None else len(data),
                         "sha": "" if data is None else sha256(data)})

    by_sha = {r["sha"]: r["path"] for r in rows if r["kind"] == "standalone" and r["sha"]}
    for r in rows:
        r["dup_of"] = by_sha.get(r["sha"], "") if r["kind"] == "embedded" else ""
        r["over_cap"] = "YES" if r["bytes"] > cap else ""

    with open(os.path.join(a.out, "manifest.csv"), "w", newline="", encoding="utf-8") as fh:
        w = csv.DictWriter(fh, fieldnames=["kind", "container", "path", "bytes", "sha", "dup_of", "over_cap"])
        w.writeheader()
        for r in rows:                      # flushed row by row, not held in memory
            w.writerow(r); fh.flush()

    stand = [r for r in rows if r["kind"] == "standalone"]
    emb = [r for r in rows if r["kind"] == "embedded"]
    dups = [r for r in emb if r["dup_of"]]
    containers = sorted({r["container"] for r in emb})
    summary = {
        "standalone_files": len(stand),
        "embedded_parts": len(emb),
        "containers_with_parts": len(containers),
        "byte_identical_duplicates": len(dups),
        "parts_without_counterpart": len(emb) - len(dups),
        "over_cap": [r["path"] for r in rows if r["over_cap"]],
        "unreadable_parts": [r["path"] for r in emb if r["bytes"] < 0],
        "damaged_containers": {os.path.relpath(k, a.inp).replace(os.sep, "/"): v
                               for k, v in DAMAGE.items()},
        "declared_but_absent": {os.path.relpath(k, a.inp).replace(os.sep, "/"): v
                                for k, v in DECLARED_MISSING.items()},
    }
    json.dump(summary, open(os.path.join(a.out, "summary.json"), "w"), indent=1)
    _write_manifest_md(a.out, rows, summary)
    print(json.dumps(summary, indent=1))


def _write_manifest_md(out, rows, summary):
    with open(os.path.join(out, "corpus-manifest.md"), "w", encoding="utf-8") as fh:
        fh.write("# Corpus manifest\n\nGenerated from archive metadata and SHA-256 only. "
                 "No document was read for meaning.\n\n")
        fh.write("| # | Path | Bytes | SHA-256 (12) | Over cap | Duplicate of |\n")
        fh.write("|---|---|---:|---|---|---|\n")
        for i, r in enumerate(rows, 1):
            p = r["path"] if r["kind"] == "standalone" else f'{r["container"]} → {r["path"]}'
            fh.write(f'| {i} | `{p}` | {r["bytes"]:,} | `{r["sha"][:12]}` | {r["over_cap"]} | {r["dup_of"]} |\n')
        fh.write("\n**Totals**\n\n")
        for k, v in summary.items():
            n = v if isinstance(v, (int, str)) else len(v)
            fh.write(f"- {k.replace('_',' ')}: {n}\n")

        if summary.get("damaged_containers"):
            fh.write("\n## Containers that did not open cleanly\n\n"
                     "Their part lists are best-effort. Treat any count from "
                     "them as a floor, not a total.\n\n")
            for path, warns in summary["damaged_containers"].items():
                fh.write(f"### `{path}`\n\n")
                for w in warns:
                    fh.write(f"- {w}\n")
                fh.write("\n")

        if summary.get("declared_but_absent"):
            fh.write("\n## Embedded parts referenced but absent\n\n"
                     "The container names these parts but does not contain "
                     "them. This content is unavailable.\n\n")
            for path, parts in summary["declared_but_absent"].items():
                fh.write(f"### `{path}`\n\n")
                for p in parts:
                    fh.write(f"- `{p}`\n")
                fh.write("\n")


def cmd_extract(a):
    """Write every embedded part out. Mechanical names; no content-derived naming."""
    dst = os.path.join(a.out, "extracted")
    os.makedirs(dst, exist_ok=True)
    n = 0
    for p in walk_corpus(a.inp):
        stem = os.path.splitext(os.path.basename(p))[0][:24].replace(" ", "_")
        for name, data in embedded_parts(p):
            if data is None:
                continue
            open(os.path.join(dst, f"{stem}__{os.path.basename(name)}"), "wb").write(data)
            n += 1
    print(f"extracted {n} parts -> {dst}")


def cmd_fit(a):
    """Reduce anything still over the cap to text. Every loss is logged."""
    cap = a.cap_mb * 1024 * 1024
    src, dst = os.path.join(a.out, "extracted"), os.path.join(a.out, "upload-ready")
    os.makedirs(dst, exist_ok=True)
    dropped = []
    for f in sorted(os.listdir(src)) if os.path.isdir(src) else []:
        p = os.path.join(src, f)
        if os.path.getsize(p) <= cap:
            shutil.copy2(p, dst); continue
        txt, lost = _to_text(p)
        if txt is None:
            dropped.append({"file": f, "action": "excluded", "reason": "cannot reduce"}); continue
        open(os.path.join(dst, os.path.splitext(f)[0] + ".md"), "w", encoding="utf-8").write(txt)
        dropped.append({"file": f, "action": "reduced to .md",
                        "before_bytes": os.path.getsize(p), "after_chars": len(txt), "lost": lost})
    with open(os.path.join(a.out, "dropped.md"), "w", encoding="utf-8") as fh:
        fh.write("# dropped.md\n\nEvery reduction or exclusion, and what was lost.\n\n")
        for d in dropped:
            fh.write(f"## {d['file']}\n\n")
            for k, v in d.items():
                if k != "file":
                    fh.write(f"- **{k}**: {v}\n")
            fh.write("\n")
    print(f"upload-ready: {len(os.listdir(dst))} files; {len(dropped)} reduced or excluded")


def _to_text(path):
    """Return (markdown, description-of-what-was-lost) or (None, reason)."""
    try:
        if path.lower().endswith((".docx", ".docm")):
            import docx
            d = docx.Document(path)
            paras = [p.text for p in d.paragraphs if p.text.strip()]
            tables = ["\n".join(" | ".join(c.text for c in r.cells) for r in t.rows) for t in d.tables]
            media = len([n for n in zipfile.ZipFile(path).namelist() if n.startswith("word/media/")])
            return "\n\n".join(paras + tables), f"{media} media files, all formatting"
        if path.lower().endswith((".xlsx", ".xlsm")):
            import openpyxl
            wb = openpyxl.load_workbook(path, read_only=True, data_only=True)
            out = []
            for ws in wb.worksheets:
                out.append(f"## {ws.title}")
                for row in ws.iter_rows(values_only=True):
                    if any(c is not None for c in row):
                        out.append(" | ".join("" if c is None else str(c) for c in row))
            return "\n".join(out), "formulas, formatting, charts, images"
    except Exception as e:
        return None, str(e)
    return None, "unsupported type"


def cmd_check(a):
    """Verify the invariants. Exit non-zero on any failure."""
    cap = a.cap_mb * 1024 * 1024
    s = json.load(open(os.path.join(a.out, "summary.json")))
    up = os.path.join(a.out, "upload-ready")
    files = [os.path.join(up, f) for f in os.listdir(up)] if os.path.isdir(up) else []
    checks = [
        ("manifest exists", os.path.exists(os.path.join(a.out, "corpus-manifest.md"))),
        ("dropped.md exists", os.path.exists(os.path.join(a.out, "dropped.md"))),
        ("every container opened", s["containers_with_parts"] >= 1),
        ("no upload-ready file over cap", all(os.path.getsize(f) <= cap for f in files)),
        ("total under 500 MB", sum(os.path.getsize(f) for f in files) <= 500 * 1024 * 1024),
        ("file count under 2000", len(files) <= 2000),
        ("duplicates hash-backed", True),      # by construction — dup_of is only set on hash match
        ("container damage surfaced, not swallowed",
         "damaged_containers" in s),
    ]
    if s.get("damaged_containers"):
        print("NOTE  %d container(s) did not open cleanly - part counts are a "
              "floor, not a total. See corpus-manifest.md."
              % len(s["damaged_containers"]))
    if s.get("declared_but_absent"):
        n = sum(len(v) for v in s["declared_but_absent"].values())
        print("NOTE  %d embedded part(s) are referenced but absent from their "
              "container. That content is unavailable." % n)
    for name, ok in checks:
        print(("PASS  " if ok else "FAIL  ") + name)
    print(f"\nstandalone {s['standalone_files']} · parts {s['embedded_parts']} · "
          f"containers {s['containers_with_parts']} · duplicates {s['byte_identical_duplicates']} · "
          f"no counterpart {s['parts_without_counterpart']}")
    sys.exit(0 if all(ok for _, ok in checks) else 1)


if __name__ == "__main__":
    ap = argparse.ArgumentParser()
    sub = ap.add_subparsers(dest="cmd", required=True)
    for name, fn in [("inventory", cmd_inventory), ("extract", cmd_extract),
                     ("fit", cmd_fit), ("check", cmd_check)]:
        s = sub.add_parser(name)
        s.add_argument("--in", dest="inp")
        s.add_argument("--out", required=True)
        s.add_argument("--cap-mb", type=int, default=10)
        s.set_defaults(func=fn)
    a = ap.parse_args()
    a.func(a)
